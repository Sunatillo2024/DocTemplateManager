import re
from docx import Document

def extract_placeholders(docx_path):
    """
    Извлекает все уникальные плейсхолдеры из .docx файла.
    Формат плейсхолдера: {{PLACEHOLDER_NAME}}
    """
    doc = Document(docx_path)
    placeholders = set()
    pattern = re.compile(r'\{\{([A-Z_]+)\}\}')

    # Обрабатываем обычные параграфы
    for para in doc.paragraphs:
        full_text = para.text
        matches = pattern.findall(full_text)
        placeholders.update(matches)

    # Обрабатываем таблицы
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    full_text = para.text
                    matches = pattern.findall(full_text)
                    placeholders.update(matches)

    return sorted(list(placeholders))


def fill_template(docx_path, data: dict, output_path: str):
    """
    Заполняет шаблон .docx данными и сохраняет в output_path.
    data: словарь вида {'FULL_NAME': 'Иван Иванов', ...}
    Реализует корректную обработку split runs.
    """
    doc = Document(docx_path)

    def replace_in_paragraph(para, data):
        # Собираем весь текст абзаца из всех runs
        full_text = ''.join(run.text for run in para.runs)

        # Заменяем плейсхолдеры в полном тексте
        for key, value in data.items():
            placeholder = f'{{{{{key}}}}}'
            if placeholder in full_text:
                full_text = full_text.replace(placeholder, str(value))

        # Если текст изменился, перезаписываем runs:
        if full_text != para.text:  # para.text даёт уже слитый текст, но лучше сравнивать с собранным
            # Очищаем все runs, кроме первого, и записываем весь текст в первый run
            if para.runs:
                para.runs[0].text = full_text
                for run in para.runs[1:]:
                    run.text = ''
            else:
                # Если runs нет (редко), создаём новый run
                para.add_run(full_text)

    # Обработка всех параграфов в теле документа
    for para in doc.paragraphs:
        replace_in_paragraph(para, data)

    # Обработка таблиц
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_paragraph(para, data)

    doc.save(output_path)